from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

NAVY   = RGBColor(0x1B, 0x35, 0x6E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF2, 0xF4, 0xF8)
MGRAY  = RGBColor(0xD0, 0xD5, 0xE0)
DKTEXT = RGBColor(0x1A, 0x1A, 0x2E)
GREEN  = RGBColor(0x1A, 0x7A, 0x4A)
AMBER  = RGBColor(0xC9, 0xA0, 0x2C)


def shade_para(para, hex_colour):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    pPr.append(shd)


def set_cell_bg(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)


def add_borders(table, colour='D0D5E0'):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), colour)
        borders.append(el)
    tblPr.append(borders)


def r(para, text, bold=False, colour=None, size=None, italic=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if colour: run.font.color.rgb = colour
    if size:   run.font.size      = Pt(size)
    return run


def sec_header(label):
    doc.add_paragraph()
    h = doc.add_paragraph()
    shade_para(h, '1B356E')
    rx = h.add_run(f'  {label}')
    rx.bold = True; rx.font.size = Pt(14); rx.font.color.rgb = WHITE


def body(text, indent=0):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for rx in p.runs:
        rx.font.size = Pt(10.5); rx.font.color.rgb = DKTEXT
    return p


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_para(cover, '1B356E')
rx = cover.add_run('\n  AA Plus Policy Training Platform\n')
rx.bold = True; rx.font.size = Pt(26); rx.font.color.rgb = WHITE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_para(sub, '1B356E')
rx2 = sub.add_run('Commercial Pricing Strategy & Rationale  |  May 2026\n')
rx2.font.size = Pt(13); rx2.font.color.rgb = RGBColor(0xC9, 0xD8, 0xFF)

sp = doc.add_paragraph(); shade_para(sp, '1B356E'); sp.add_run(' ')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — POSITIONING
# ══════════════════════════════════════════════════════════════════════════════
sec_header('1.  Product Positioning')

body(
    "AA Plus Policy Training Platform is a managed India-native compliance platform. "
    "It bundles 11+ pre-built regulatory training modules with auto-generated verifiable "
    "certificates, multi-tenant organisation management, department-level access control, "
    "audit trails, and Excel-based compliance reporting. "
    "AA Plus manages the platform; clients log in, assign training, and track compliance."
)
body(
    "Unlike self-serve LMS tools, clients do not configure or maintain the platform. "
    "AA Plus handles module updates as Indian regulations evolve — ensuring organisations "
    "always train on current, legally accurate content."
)

for label, val in [
    ("Closest global comparison:", "Ethena (US) — $50/user/year for their full catalog, no India-specific content"),
    ("India POSH-only comparison:", "POSH Check — Rs.55,000/year for 100 users, POSH module only"),
    ("Full LMS comparison:",        "Disprz — ~Rs.250/user/month, enterprise-only, no bundled India regulatory content"),
]:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.left_indent = Inches(0.3)
    r(bp, label + '  ', bold=True, colour=NAVY, size=10.5)
    r(bp, val, size=10.5, colour=DKTEXT)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PRICING PHILOSOPHY
# ══════════════════════════════════════════════════════════════════════════════
sec_header('2.  Pricing Philosophy')

body(
    "Pricing is based on the value and risk protection delivered — not on undercutting "
    "cheaper or inferior products. The platform eliminates significant legal and financial "
    "risk for every customer:"
)

risks = [
    ("POSH Act violation:",    "Up to Rs.50,000 fine + criminal liability for ICC non-compliance"),
    ("DPDP Act breach:",       "Up to Rs.250 crore penalty under the Digital Personal Data Protection Act 2023"),
    ("CERT-In non-reporting:", "Up to Rs.1 crore penalty for failure to report cyber incidents within 6 hours"),
    ("Anti-Corruption:",       "Criminal prosecution and licence risk for listed companies and MNC subsidiaries"),
]
rt = doc.add_table(rows=len(risks), cols=2)
rt.style = 'Table Grid'
add_borders(rt)
for i, (lbl, val) in enumerate(risks):
    bg = 'FFF8E7' if i % 2 == 0 else 'FFFDF0'
    set_cell_bg(rt.rows[i].cells[0], bg)
    set_cell_bg(rt.rows[i].cells[1], bg)
    r(rt.rows[i].cells[0].paragraphs[0], lbl, bold=True, colour=AMBER, size=10.5)
    r(rt.rows[i].cells[1].paragraphs[0], val, size=10.5, colour=DKTEXT)
rt.columns[0].width = Cm(5.5)
rt.columns[1].width = Cm(13)

for title, detail in [
    ("Value-based pricing signals credibility",
     "Compliance buyers (Legal, HR, CFO) associate low prices with low rigour. "
     "A Rs.90,000 contract positions AA Plus as a serious compliance partner."),
    ("Managed platform commands a premium",
     "AA Plus updates modules as laws change, manages the platform backend, and provides "
     "onboarding support. This is a service, not just software."),
    ("Margin funds content quality",
     "Indian compliance law changes frequently. Adequate pricing funds timely module "
     "updates, legal review, and new modules as regulations emerge."),
    ("Published INR pricing builds trust",
     "Prices are published in INR with GST. Indian CFOs require GST-compliant INR invoices "
     "to claim 18% input credit. Hides nothing; no sales call just to get a price."),
]:
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.paragraph_format.left_indent = Inches(0.2)
    r(tp, '  ' + title, bold=True, colour=NAVY, size=11)
    dp = doc.add_paragraph(detail)
    dp.paragraph_format.left_indent = Inches(0.5)
    dp.paragraph_format.space_after = Pt(2)
    for rx in dp.runs:
        rx.font.size = Pt(10.5); rx.font.color.rgb = DKTEXT


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SUBSCRIPTION PLANS
# ══════════════════════════════════════════════════════════════════════════════
sec_header('3.  Subscription Plans')

# NOTE: only features that are actually built in the webapp are listed here.
# Admin seat limits, custom module upload, parent/child orgs, and automated
# billing are deferred to a post-launch build.

tiers = [
    {
        "name":     "ESSENTIALS",
        "colour":   "C9A02C",
        "employees":"Up to 25 employees",
        "price":    "Rs.36,000 / year  (+18% GST)",
        "per_emp":  "Rs.1,440 / employee / year",
        "per_day":  "Rs.3.95 / employee / day",
        "best_for": "Startups, boutique firms, small offices",
        "includes": [
            "All 11+ compliance modules — POSH Act, Cybersecurity/CERT-In, DPDP Act 2023, Anti-Corruption, Code of Conduct, Whistleblower Policy, HSE, Company Law, Information Security, AI Policy (3-tier)",
            "Digital certificates with QR code public verification link",
            "Admin dashboard with user and module management",
            "Learner dashboard with interactive quiz engine and per-section progress tracking",
            "Email invites via secure token system (individual and bulk CSV upload)",
            "Compliance report — module completion status per user",
            "AA Plus manages all platform updates and content revisions as regulations change",
        ],
    },
    {
        "name":     "BUSINESS",
        "colour":   "1B6B3A",
        "employees":"Up to 100 employees",
        "price":    "Rs.90,000 / year  (+18% GST)",
        "per_emp":  "Rs.900 / employee / year",
        "per_day":  "Rs.2.47 / employee / day",
        "best_for": "Growing companies, multi-team organisations",
        "includes": [
            "Everything in Essentials",
            "Department creation and management — segment users by team or function",
            "Manager dashboard — team-level compliance view showing each member's status",
            "Excel export — compliance reports and full activity audit trail",
            "Automated email nudges tied to custom per-module deadlines",
            "Module due-date tracking with overdue alerts",
            "Dedicated onboarding call with the AA Plus team",
        ],
    },
    {
        "name":     "PROFESSIONAL",
        "colour":   "1B356E",
        "employees":"Up to 250 employees",
        "price":    "Rs.1,80,000 / year  (+18% GST)",
        "per_emp":  "Rs.720 / employee / year",
        "per_day":  "Rs.1.97 / employee / day",
        "best_for": "Listed companies, multi-department organisations",
        "includes": [
            "Everything in Business",
            "POSH Policy Wizard — configure ICC member details, effective date, and contact information; generates downloadable POSH policy PDF",
            "Slack and MS Teams webhook notifications triggered on module completion",
            "RPT (Related Party Transaction) Simulator — interactive decision-making tool with department-level access control",
            "AI Policy Tier module — 3-tier governance framework; org admin completes a quiz to auto-assign the correct tier",
            "Organisation logo white-labelling on certificates and invitation emails",
            "Leaderboard — top performers ranked by completion and quiz scores",
            "Full activity audit log — event-level trail with date-range filtering and Excel export",
        ],
    },
    {
        "name":     "CORPORATE",
        "colour":   "6B1B4A",
        "employees":"Up to 500 employees",
        "price":    "Rs.3,00,000 / year  (+18% GST)",
        "per_emp":  "Rs.600 / employee / year",
        "per_day":  "Rs.1.64 / employee / day",
        "best_for": "Large companies, BFSI / IT / pharma, listed entities",
        "includes": [
            "Everything in Professional",
            "Department-level module assignment — assign specific modules to specific teams only",
            "Bulk certificate export — download all completion certificates for a module in one action",
            "Priority support with 24-hour email response SLA",
            "Custom support email address configured per organisation",
            "Price locked for 2 years — no increases at renewal",
            "Annual compliance health review call with the AA Plus team",
        ],
    },
    {
        "name":     "ENTERPRISE",
        "colour":   "2C4E8A",
        "employees":"500+ employees",
        "price":    "Custom quote",
        "per_emp":  "~Rs.350-500 / employee / year",
        "per_day":  "Rs.0.96-1.37 / employee / day",
        "best_for": "Large enterprises, conglomerates, heavily regulated sectors",
        "includes": [
            "Everything in Corporate",
            "Dedicated Customer Success Manager from the AA Plus team",
            "Custom SLA agreement",
            "SSO integration via SAML / OIDC — roadmap",
            "HRMS integration — Darwinbox, Keka, Zoho People — roadmap",
            "Volume discount for multi-year contracts — 15% off (2-year), 20% off (3-year)",
            "White-glove onboarding — AA Plus configures the full org, departments, and module assignments",
        ],
    },
]

TIER_HL = ['FFF8E7','EAF4F0','E8EEF9','F5E8F5','E8F0FF']

for i, tier in enumerate(tiers):
    doc.add_paragraph()
    chex = tier['colour']
    crgb = RGBColor(int(chex[0:2],16), int(chex[2:4],16), int(chex[4:6],16))

    th = doc.add_paragraph()
    shade_para(th, chex)
    rx = th.add_run(f"  {tier['name']}  —  {tier['employees']}  |  {tier['price']}")
    rx.bold = True; rx.font.size = Pt(13); rx.font.color.rgb = WHITE

    kt = doc.add_table(rows=1, cols=3)
    kt.style = 'Table Grid'
    add_borders(kt)
    for j, (lbl, val) in enumerate([
        ('Per employee / year', tier['per_emp']),
        ('Per employee / day',  tier['per_day']),
        ('Best for',            tier['best_for']),
    ]):
        set_cell_bg(kt.rows[0].cells[j], 'F2F4F8')
        r(kt.rows[0].cells[j].paragraphs[0], lbl + ':  ', bold=True, colour=crgb, size=10)
        r(kt.rows[0].cells[j].paragraphs[0], val, size=10, colour=DKTEXT)

    il = doc.add_paragraph()
    il.paragraph_format.left_indent = Inches(0.15)
    il.paragraph_format.space_before = Pt(6)
    r(il, "What's included:", bold=True, colour=NAVY, size=10.5)

    for item in tier['includes']:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent = Inches(0.5)
        bp.paragraph_format.space_after = Pt(1)
        r(bp, item, size=10, colour=DKTEXT)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — ADD-ONS
# ══════════════════════════════════════════════════════════════════════════════
sec_header('4.  Optional Module Add-Ons')

body(
    "For organisations that need a specific module rather than the full suite, "
    "or want to extend beyond their base plan:"
)

addons = [
    ["POSH-Only Package  (any org size, unlimited users)", "Rs.45,000 / year",       "ICC wizard, POSH certificates, admin reports — full POSH Act compliance"],
    ["AI Policy Module (3-Tier Framework)",               "Rs.15,000 / year",       "Add to any base plan; org admin quiz auto-assigns the correct tier"],
    ["RPT Simulator",                                     "Rs.20,000 / year",       "Interactive Related Party Transaction decision-making tool"],
    ["Extra Admin Seat",                                  "Rs.5,000 / seat / year", "Additional admin accounts beyond standard allocation"],
]

at = doc.add_table(rows=1, cols=3)
at.style = 'Table Grid'
add_borders(at)
for j, hdr in enumerate(['Add-On', 'Annual Price', 'Notes']):
    set_cell_bg(at.rows[0].cells[j], '1B356E')
    r(at.rows[0].cells[j].paragraphs[0], hdr, bold=True, colour=WHITE, size=10.5)

for i, rd in enumerate(addons):
    row = at.add_row()
    bg = 'F2F4F8' if i % 2 == 0 else 'FFFFFF'
    for j, val in enumerate(rd):
        set_cell_bg(row.cells[j], bg)
        r(row.cells[j].paragraphs[0], val, bold=(j==1), colour=NAVY if j==1 else DKTEXT, size=10)

at.columns[0].width = Cm(6.5)
at.columns[1].width = Cm(4.5)
at.columns[2].width = Cm(7.5)

doc.add_paragraph()
sn = doc.add_paragraph()
sn.paragraph_format.left_indent = Inches(0.2)
r(sn, 'Note on POSH-Only: ', bold=True, colour=NAVY, size=10.5)
r(sn,
    'The POSH-Only package at Rs.45,000 is positioned above POSH Check (Rs.30,000-55,000) '
    'because it includes a certificate system, activity audit trail, and manager reporting '
    'that POSH Check does not offer — not as a discount entry point but as a premium alternative.',
    size=10.5, colour=DKTEXT)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — VOLUME PRICING
# ══════════════════════════════════════════════════════════════════════════════
sec_header('5.  Volume & Enterprise Pricing')

vt = doc.add_table(rows=1, cols=3)
vt.style = 'Table Grid'
add_borders(vt)
for j, hdr in enumerate(['Employee Count', 'Per Employee / Year', 'Estimated Annual Cost']):
    set_cell_bg(vt.rows[0].cells[j], '1B356E')
    r(vt.rows[0].cells[j].paragraphs[0], hdr, bold=True, colour=WHITE, size=10.5)

for i, rd in enumerate([
    ['500-999',     'Rs.500',     'Rs.2,50,000 - Rs.4,99,000'],
    ['1,000-2,499', 'Rs.400',     'Rs.4,00,000 - Rs.9,99,000'],
    ['2,500-4,999', 'Rs.300',     'Rs.7,50,000 - Rs.14,99,000'],
    ['5,000+',      'Rs.200-250', 'Custom quote'],
]):
    row = vt.add_row()
    bg = 'F2F4F8' if i % 2 == 0 else 'FFFFFF'
    for j, val in enumerate(rd):
        set_cell_bg(row.cells[j], bg)
        r(row.cells[j].paragraphs[0], val, bold=(j==1), colour=NAVY if j==1 else DKTEXT, size=10.5)

doc.add_paragraph()
dp = doc.add_paragraph()
dp.paragraph_format.left_indent = Inches(0.2)
r(dp, 'Multi-year discounts:  ', bold=True, colour=NAVY, size=10.5)
r(dp, '2-year upfront — 10% off  |  3-year upfront — 20% off  |  Price locked for contract duration', size=10.5, colour=DKTEXT)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — COMPETITOR BENCHMARKS
# ══════════════════════════════════════════════════════════════════════════════
sec_header('6.  Competitor Pricing Benchmarks')

body('Annual cost comparison for 100 employees (converted to INR at Rs.84/USD). AA Plus highlighted in green.')

ct = doc.add_table(rows=1, cols=4)
ct.style = 'Table Grid'
add_borders(ct)
for j, hdr in enumerate(['Platform', 'What You Get', 'Annual Cost (100 users)', 'Notes']):
    set_cell_bg(ct.rows[0].cells[j], '1B356E')
    r(ct.rows[0].cells[j].paragraphs[0], hdr, bold=True, colour=WHITE, size=10)

for i, rd in enumerate([
    ['AA Plus — Business',     '11+ India compliance modules + certs + admin + audit', 'Rs.90,000',      'Full managed platform'],
    ['POSH Check Pro (India)', 'POSH training only',                                   'Rs.55,000',      'POSH Act only; no other modules'],
    ['Ethena (US)',            'Harassment prevention only',                           '~Rs.1,70,000',   '$20/user/yr; no India content'],
    ['Ethena Full Catalog',    '250+ topics — US-focused',                            '~Rs.4,20,000',   '$50/user/yr; no India regulatory modules'],
    ['EasyLlama (US)',         '30+ compliance topics',                                '~Rs.1,68,000',   '$19.95/seat/yr; minimum 10 seats'],
    ['Disprz (India)',         'Full LMS + some India content',                        '~Rs.3,00,000',   '~$3/user/month; enterprise-only'],
    ['TalentLMS Grow',         'LMS shell only — no compliance content',               '~Rs.2,30,000',   '$229/month; customer must source all content'],
    ['KnowBe4',                'Cybersecurity awareness only',                         '~Rs.3,40,000',   '~$3.19/seat/month; single topic'],
    ['SAP Litmos + content',   'LMS + basic compliance bundle',                        '~Rs.10,00,000+', 'Custom quote; enterprise minimum'],
]):
    row = ct.add_row()
    bg = 'DFF2E8' if i == 0 else ('F2F4F8' if i % 2 == 0 else 'FFFFFF')
    for j, val in enumerate(rd):
        set_cell_bg(row.cells[j], bg)
        bold = (i == 0)
        colour = GREEN if (i == 0 and j in (0,2)) else DKTEXT
        r(row.cells[j].paragraphs[0], val, bold=bold, colour=colour, size=10)

ct.columns[0].width = Cm(3.8)
ct.columns[1].width = Cm(5.5)
ct.columns[2].width = Cm(3.8)
ct.columns[3].width = Cm(5.4)

doc.add_paragraph()
kp = doc.add_paragraph()
kp.paragraph_format.left_indent = Inches(0.2)
r(kp, 'Key insight: ', bold=True, colour=NAVY, size=10.5)
r(kp,
    "At Rs.90,000/year, the Business plan delivers 11 India-specific regulatory modules "
    "for less than what Ethena charges for harassment prevention alone (~Rs.1,70,000). "
    "Compared to Ethena's full catalog (~Rs.4,20,000 for 100 users), AA Plus delivers "
    "equivalent or greater value at less than one-quarter the price — with content tailored to Indian law.",
    size=10.5, colour=DKTEXT)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — BILLING & PAYMENT
# ══════════════════════════════════════════════════════════════════════════════
sec_header('7.  Billing & Payment')

body(
    "AA Plus operates on a high-touch, relationship-first billing model. "
    "There is no self-serve checkout — organisations sign up via a pilot request or direct contact, "
    "and AA Plus activates the account upon payment confirmation."
)

for title, detail in [
    ('Pilot / trial',
     'Contact praveen@aaplusconsultants.com to request a free 14-day pilot for up to 10 users. '
     'Full-featured access. No credit card required.'),
    ('Annual billing (default)',
     'AA Plus sends a GST-compliant INR invoice via email. '
     'Payment via bank transfer (NEFT/RTGS), UPI, or cheque. '
     'Account is activated within 1 business day of payment confirmation.'),
    ('Monthly billing (+20% premium)',
     'Available for plans under 100 employees on request. '
     'Invoice sent at the start of each month.'),
    ('GST',
     '18% GST applicable on all plans. HSN code 998314 (SaaS). '
     'GSTIN included on all invoices for input credit claim.'),
    ('Multi-year discounts',
     '10% off for 2-year upfront. 20% off for 3-year upfront. '
     'Price locked for the full contract duration.'),
    ('Renewal',
     '30-day renewal reminder sent to the org Admin. '
     'AA Plus sends an annual compliance summary (completion rates by module and department) '
     'at the end of each subscription year.'),
]:
    bp = doc.add_paragraph()
    bp.paragraph_format.left_indent = Inches(0.2)
    bp.paragraph_format.space_before = Pt(5)
    r(bp, '  ' + title + ':  ', bold=True, colour=NAVY, size=11)
    r(bp, detail, size=10.5, colour=DKTEXT)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
sec_header('8.  Quick Reference — Plan Summary')

st = doc.add_table(rows=1, cols=5)
st.style = 'Table Grid'
add_borders(st)
for j, hdr in enumerate(['Plan', 'Employees', 'Annual Price', 'Rs./Employee/Yr', 'Key Additions']):
    set_cell_bg(st.rows[0].cells[j], '1B356E')
    r(st.rows[0].cells[j].paragraphs[0], hdr, bold=True, colour=WHITE, size=10)

TIER_BG = ['FFF8E7','EAF4F0','E8EEF9','F5E8F5','E8F0FF']
for i, rd in enumerate([
    ['Essentials',   'Up to 25',  'Rs.36,000',   'Rs.1,440', 'All modules, certificates, reports'],
    ['Business',     'Up to 100', 'Rs.90,000',   'Rs.900',   '+ Departments, Excel export, nudges, manager dashboard, onboarding call'],
    ['Professional', 'Up to 250', 'Rs.1,80,000', 'Rs.720',   '+ POSH wizard, Slack/Teams, RPT Simulator, AI Policy, leaderboard, audit log'],
    ['Corporate',    'Up to 500', 'Rs.3,00,000', 'Rs.600',   '+ Dept-level modules, bulk cert export, priority SLA, price lock, annual review call'],
    ['Enterprise',   '500+',      'Custom',       'Rs.350-500','+ Dedicated CSM, SSO/HRMS (roadmap), white-glove onboarding, volume discounts'],
]):
    row = st.add_row()
    for j, val in enumerate(rd):
        set_cell_bg(row.cells[j], TIER_BG[i])
        r(row.cells[j].paragraphs[0], val, bold=(j in (0,2,3)), colour=NAVY if j in (0,2,3) else DKTEXT, size=10)


# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_para(foot, 'F2F4F8')
r(foot,
  'AA Plus Consultants  |  praveen@aaplusconsultants.com  |  Confidential — May 2026',
  size=9, colour=MGRAY)

out = r'C:\Users\pc\Desktop\PolicyTraining\AA_Plus_Pricing_Strategy.docx'
doc.save(out)
print(f'Saved: {out}')
